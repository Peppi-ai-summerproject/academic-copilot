"""Document Loader for RAG Knowledge Base - Issue #57"""

from pathlib import Path
from langchain_core.documents import Document
from docx import Document as DocxDocument
from pypdf import PdfReader


class DocumentLoader:

    def load_docx(self, file_path):
        doc = DocxDocument(file_path)
        text = '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
        return Document(page_content=text, metadata={"source": str(file_path), "filename": Path(file_path).name, "file_type": "docx"})

    def load_pdf(self, file_path):
        reader = PdfReader(file_path)
        text = '\n'.join([page.extract_text() for page in reader.pages if page.extract_text()])
        return Document(page_content=text, metadata={"source": str(file_path), "filename": Path(file_path).name, "file_type": "pdf"})

    def load_directory(self, directory):
        documents = []
        for file_path in Path(directory).iterdir():
            if file_path.suffix.lower() == ".docx":
                print(f"Loading: {file_path.name}")
                documents.append(self.load_docx(str(file_path)))
            elif file_path.suffix.lower() == ".pdf":
                print(f"Loading: {file_path.name}")
                documents.append(self.load_pdf(str(file_path)))
        print(f"Total loaded: {len(documents)} documents")
        return documents
